from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt

def create_practical_training_log_book(department, student_name, reg_no, company, week_no, from_date, to_date, data_dictionary):
    """
    data_dictionary = {
        monday:{
            'date': date,
            'activity': activity
        },
        tuesday:{
            'date': date,
            'activity': activity
        }
    }
    """
    # Create a new Document
    doc = Document()

    # Title
    title = doc.add_heading('UNIVERSITY OF DAR ES SALAAM', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # College and Department
    doc.add_heading('COLLEGE OF INFORMATION AND COMMUNICATION TECHNOLOGIES', level=2)
    department = doc.add_heading(f'DEPARTMENT OF {department}', level=2)
    department.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Practical Training Log - Book
    doc.add_heading('PRACTICAL TRAINING LOG – BOOK', level=1)
    doc.add_paragraph()

    # Student Information
    student_table = doc.add_table(rows=1, cols=2)
    student_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set column widths (50% for each column)
    student_table.columns[0].width = Inches(4)
    student_table.columns[1].width = Inches(2)

    # Add text to the cells
    cell_0_0 = student_table.cell(0, 0)
    cell_0_1 = student_table.cell(0, 1)

    cell_0_0.text = f'STUDENTS NAME: {student_name}'
    cell_0_1.text = f'REG. NO: {reg_no}'

    # add another row with a single column for company name
    company_row = student_table.add_row()
    company_cell = company_row.cells[0]
    company_cell.text = f'COMPANY/INSTITUTION: {company}'
    company_cell.merge(student_table.cell(1, 1)) 

    week_log_table = doc.add_table(rows=1, cols=3)
    week_log_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    week_number_cell = week_log_table.cell(0, 0)
    week_number_cell.text = f'WEEK NO: {week_no}'
    from_date_cell = week_log_table.cell(0, 1)
    from_date_cell.text = f'FROM: {from_date}'
    to_date_cell = week_log_table.cell(0,2)
    to_date_cell.text = f'TO: {to_date}'

    # set cell widths 30% each
    week_log_table.columns[0].width = Inches(1.5)
    week_log_table.columns[1].width = Inches(2.25)
    week_log_table.columns[2].width = Inches(2.25)

    # add an empty paragraph with 2 new line breaks
    doc.add_paragraph('\n')

    # Days and Activities
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
    date = '17/07/2023'  # Replace with the actual date

    # Create a table for Days and Activities
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table.autofit = False

    # Set vertical alignment to top for all cells
    for cell in table._cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Set column widths
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(4)

    # Add Days and Activities column headers
    table.cell(0, 0).text = 'DAY / DATE'
    table.cell(0, 1).text = 'ACTIVITY'

    # Populate the table with Days, Dates, and an empty cell for activities
    for i, day in enumerate(days):
        table.cell(i + 1, 0).text = f" {day} \n \n {data_dictionary[day]['date']}"
        table.cell(i + 1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.cell(i + 1, 1).text = data_dictionary[day]['activity']

    # Details of the Main Job of the Week
    doc.add_heading('Details Of the Main Job of the Week', level=2)

    # Operation and Machinery/Tools Used
    doc.add_heading('Operation:', level=3)
    machinery_heading = doc.add_heading('Machinery/Tools Used', level=3)

    # Add a table for machinery/tools used
    machinery_table = doc.add_table(rows=7, cols=2)
    machinery_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for row in machinery_table.rows:
        for cell in row.cells:
            cell.text = ' '

    # Comments from Industrial Supervisor
    doc.add_heading('Comments from Industrial Supervisor', level=2)

    # Name and Signature
    name_paragraph = doc.add_paragraph('Name: ............................................................  Signature: .......................................')
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Detailed Diagram of the Main Job
    doc.add_heading('Detailed Diagram of the Main Job', level=2)
    diagram_paragraph = doc.add_paragraph()
    diagram_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Save the Document
    doc.save('docs/Practical_Training_Log_Book.docx')

if __name__ == '__main__':

    data_dictionary = {
        'Monday':{
            'date': '24/04/2023',
            'activity': 'Lorem Impsum template activity for logging daily activity, i want to put a date here but this is a dictionary that i am going to store data here'
        },
        'Tuesday':{
            'date': '24/04/2023',
            'activity': 'Lorem Impsum template activity for logging daily activity, i want to put a date here but this is a dictionary that i am going to store data here'
        },
        'Wednesday':{
            'date': '24/04/2023',
            'activity': 'Lorem Impsum template activity for logging daily activity, i want to put a date here but this is a dictionary that i am going to store data here'
        },
        'Thursday':{
            'date': '24/04/2023',
            'activity': 'Lorem Impsum template activity for logging daily activity, i want to put a date here but this is a dictionary that i am going to store data here'
        },
        'Friday':{
            'date': '24/04/2023',
            'activity': 'Lorem Impsum template activity for logging daily activity, i want to put a date here but this is a dictionary that i am going to store data here'
        }
    }

    
    department = "COMPUTER ENGINEERING" 
    student_name = "GULAY, EDGAR EDWARD" 
    reg_no = "2021-04-02100" 
    company = "UDICTI"
    week_no = 1 
    from_date = "17/07/2021"
    to_date = "23/08/2021"

    # create_practical_training_log_book()
    create_practical_training_log_book(department, student_name, reg_no, company, week_no, from_date, to_date, data_dictionary)


# create deliver make it sustainable